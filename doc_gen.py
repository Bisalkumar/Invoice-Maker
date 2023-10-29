from docx import Document

# Load the document
doc = Document("invoice_template.docx")

# Sample data
invoice_list = [[2, "pen", 0.5, 1],
                [1, "paper pack", 5, 5],
                [2, "notebook", 2, 4]]

data = {
    "name": "john",
    "phone": "555-55555",
    "subtotal": 10,
    "salestax": "10%",
    "total": 9
}

# Replace placeholders in the document
# Assuming that your placeholders are in a format like {{name}}, {{phone}}, etc.
for paragraph in doc.paragraphs:
    for key, value in data.items():
        if '{{' + key + '}}' in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace('{{' + key + '}}', str(value))

# Assuming the invoice list goes into a table
table = doc.tables[0]  # Assuming the first table is where the invoice list should go

for item in invoice_list:
    cells = table.add_row().cells  # add a new row and get the cells
    for index, value in enumerate(item):
        cells[index].text = str(value)

# Save the modified document
doc.save("new_invoice.docx")
